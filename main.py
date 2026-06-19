"""
AI Resume Analyzer - FastAPI + Groq
Analyzes resumes (PDF/DOCX/TXT) against an optional job description
using Groq's LLM API, and returns structured feedback as JSON.
"""

import os
import json
import logging
from typing import Optional

from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from groq import Groq, APIError, APIConnectionError, RateLimitError

import pdfplumber
import docx  # python-docx

# --------------------------------------------------------------------------
# Configuration
# --------------------------------------------------------------------------

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("resume_analyzer")

GROQ_API_KEY = os.environ.get("GROQ_API_KEY")
GROQ_MODEL = os.environ.get("GROQ_MODEL", "llama-3.3-70b-versatile")

MAX_FILE_SIZE_MB = 5
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

app = FastAPI(
    title="AI Resume Analyzer",
    description="Upload a resume and get instant AI-powered feedback using Groq.",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Lazily-created Groq client so the app can still start (and serve the
# frontend / health check) even if the API key isn't set yet.
_groq_client: Optional[Groq] = None


def get_groq_client() -> Groq:
    global _groq_client
    if not GROQ_API_KEY:
        raise HTTPException(
            status_code=500,
            detail=(
                "GROQ_API_KEY is not set on the server. "
                "Set it as an environment variable before starting the app."
            ),
        )
    if _groq_client is None:
        _groq_client = Groq(api_key=GROQ_API_KEY)
    return _groq_client


# --------------------------------------------------------------------------
# Text extraction helpers
# --------------------------------------------------------------------------

def extract_text_from_pdf(file_bytes: bytes) -> str:
    import io

    text_chunks = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_chunks.append(page_text)
    except Exception as exc:
        logger.exception("Failed to parse PDF")
        raise HTTPException(status_code=422, detail=f"Could not read PDF file: {exc}")

    text = "\n".join(text_chunks).strip()
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No extractable text found in PDF. It may be a scanned image without OCR.",
        )
    return text


def extract_text_from_docx(file_bytes: bytes) -> str:
    import io

    try:
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs]
        # Also pull text out of any tables (common in resumes)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        paragraphs.append(cell.text)
        text = "\n".join(p for p in paragraphs if p.strip())
    except Exception as exc:
        logger.exception("Failed to parse DOCX")
        raise HTTPException(status_code=422, detail=f"Could not read DOCX file: {exc}")

    if not text.strip():
        raise HTTPException(status_code=422, detail="No extractable text found in DOCX file.")
    return text


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8", errors="ignore")
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Could not read text file: {exc}")
    if not text.strip():
        raise HTTPException(status_code=422, detail="The uploaded text file is empty.")
    return text


def extract_resume_text(filename: str, file_bytes: bytes) -> str:
    ext = os.path.splitext(filename.lower())[1]
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed types: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )


# --------------------------------------------------------------------------
# Prompt construction & Groq call
# --------------------------------------------------------------------------

ANALYSIS_SYSTEM_PROMPT = """You are an expert technical recruiter and resume coach with 15 years \
of experience reviewing resumes across tech, finance, and consulting. \
You give precise, actionable, honest feedback. You always respond with \
strictly valid JSON only - no markdown fences, no commentary outside the JSON object."""


def build_user_prompt(resume_text: str, job_description: Optional[str]) -> str:
    jd_block = (
        f"\nThe candidate is targeting this specific job description:\n\"\"\"\n{job_description.strip()}\n\"\"\"\n"
        "Evaluate the resume's fit against this job description specifically, including a match score."
        if job_description and job_description.strip()
        else "\nNo specific job description was provided. Evaluate the resume generally for quality, "
        "clarity, and impact for a competitive job search in its apparent field."
    )

    return f"""Analyze the following resume.
{jd_block}

Resume text:
\"\"\"
{resume_text.strip()[:12000]}
\"\"\"

Respond with ONLY a JSON object (no markdown, no backticks) matching exactly this schema:

{{
  "overall_score": <integer 0-100, overall resume quality/fit score>,
  "summary": "<2-3 sentence overall assessment>",
  "strengths": ["<short bullet>", "..."],
  "weaknesses": ["<short bullet>", "..."],
  "missing_keywords": ["<keyword or skill missing relative to the role/job description>", "..."],
  "ats_friendliness": {{
      "score": <integer 0-100>,
      "notes": "<1-2 sentences on formatting/ATS parseability>"
  }},
  "section_feedback": {{
      "contact_info": "<short feedback or 'OK'>",
      "summary_or_objective": "<short feedback or 'OK'>",
      "experience": "<short feedback or 'OK'>",
      "education": "<short feedback or 'OK'>",
      "skills": "<short feedback or 'OK'>"
  }},
  "job_match_score": <integer 0-100 or null if no job description was given>,
  "suggested_improvements": ["<concrete, specific, actionable suggestion>", "..."],
  "rewritten_bullet_examples": [
      {{"original": "<a weak bullet point pulled from the resume, or empty string if none found>",
        "improved": "<a stronger rewritten version using action verbs and quantified impact>"}}
  ]
}}

Rules:
- strengths, weaknesses, missing_keywords, suggested_improvements: 3-6 items each.
- rewritten_bullet_examples: 2-4 items, pulled from real lines in the resume when possible.
- Be specific and reference actual content from the resume, not generic advice.
- Output ONLY the JSON object."""


def call_groq_analysis(resume_text: str, job_description: Optional[str]) -> dict:
    client = get_groq_client()
    user_prompt = build_user_prompt(resume_text, job_description)

    try:
        completion = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": ANALYSIS_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.4,
            max_tokens=2000,
            response_format={"type": "json_object"},
        )
    except RateLimitError as exc:
        logger.warning("Groq rate limit hit: %s", exc)
        raise HTTPException(status_code=429, detail="Groq API rate limit exceeded. Please try again shortly.")
    except APIConnectionError as exc:
        logger.error("Groq connection error: %s", exc)
        raise HTTPException(status_code=502, detail="Could not connect to Groq API. Check network/API key.")
    except APIError as exc:
        logger.error("Groq API error: %s", exc)
        raise HTTPException(status_code=502, detail=f"Groq API returned an error: {exc}")
    except Exception as exc:
        logger.exception("Unexpected error calling Groq")
        raise HTTPException(status_code=500, detail=f"Unexpected error calling Groq API: {exc}")

    if not completion.choices:
        raise HTTPException(status_code=502, detail="Groq API returned no choices.")

    raw_content = completion.choices[0].message.content
    if not raw_content:
        raise HTTPException(status_code=502, detail="Groq API returned an empty response.")

    cleaned = raw_content.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:].strip()

    try:
        result = json.loads(cleaned)
    except json.JSONDecodeError as exc:
        logger.error("Failed to parse Groq JSON output: %s\nRaw: %s", exc, raw_content)
        raise HTTPException(
            status_code=502,
            detail="Groq API returned malformed JSON. Please try again.",
        )

    return result


# --------------------------------------------------------------------------
# API routes
# --------------------------------------------------------------------------

class HealthResponse(BaseModel):
    status: str
    groq_configured: bool
    model: str


@app.get("/api/health", response_model=HealthResponse)
def health_check():
    return HealthResponse(
        status="ok",
        groq_configured=bool(GROQ_API_KEY),
        model=GROQ_MODEL,
    )


@app.post("/api/analyze")
async def analyze_resume(
    file: UploadFile = File(..., description="Resume file: .pdf, .docx, or .txt"),
    job_description: Optional[str] = Form(None, description="Optional job description to match against"),
):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file was uploaded.")

    ext = os.path.splitext(file.filename.lower())[1]
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{ext}'. Allowed types: {', '.join(sorted(ALLOWED_EXTENSIONS))}",
        )

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum allowed size is {MAX_FILE_SIZE_MB}MB.",
        )

    resume_text = extract_resume_text(file.filename, file_bytes)

    if len(resume_text.strip()) < 50:
        raise HTTPException(
            status_code=422,
            detail="Extracted resume text is too short to analyze meaningfully (less than 50 characters).",
        )

    analysis = call_groq_analysis(resume_text, job_description)

    return JSONResponse(
        content={
            "filename": file.filename,
            "characters_extracted": len(resume_text),
            "model_used": GROQ_MODEL,
            "analysis": analysis,
        }
    )


@app.get("/")
def serve_frontend():
    index_path = os.path.join(os.path.dirname(__file__), "static", "index.html")
    if not os.path.exists(index_path):
        raise HTTPException(status_code=404, detail="Frontend not found.")
    return FileResponse(index_path)


# Mount static assets (css/js) AFTER routes so "/" route above takes priority
app.mount("/static", StaticFiles(directory=os.path.join(os.path.dirname(__file__), "static")), name="static")


# --------------------------------------------------------------------------
# Global exception handlers (so the API never leaks raw tracebacks)
# --------------------------------------------------------------------------

@app.exception_handler(Exception)
async def unhandled_exception_handler(request, exc):
    logger.exception("Unhandled exception")
    return JSONResponse(
        status_code=500,
        content={"detail": f"Internal server error: {str(exc)}"},
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
